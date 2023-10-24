from docx import Document as d
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import assitent_info #Просто обращаемся к тексту
import streamlit as st
import pandas as pd
from import_google_table import table_vector

# doc = d('C:\\Users\\kushhov\\Desktop\\vector-main\\template.docx')
def fulfil_temp(cblank,type_scheme): # Сама функция вывода всего и вся в test.doc
    doc = d('template.docx')
    #doc = d('./template.docx')
    # Таблица шапки
    doc.paragraphs[2].text = f"Узел регулирующий {cblank['vector']}"
    doc.tables[0].rows[0].cells[1].text = cblank['orderer']
    doc.tables[0].rows[0].cells[3].text = 'общепромышленное'
    doc.tables[0].rows[1].cells[1].text = cblank['object']
    doc.tables[0].rows[2].cells[1].text = cblank['manager']
    doc.tables[0].rows[2].cells[3].text = cblank['developer']
    doc.tables[0].autofit
    #st.write(doc.tables[0].style)
    doc.tables[2].rows[1].cells[1].text = f"На основании {int(cblank['glycol'])}% {cblank['glycol type'][:-1]}я" if cblank['glycol'] else assitent_info.pure_water
    doc.tables[2].rows[1].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.tables[2].rows[3].cells[1].text = str(cblank['consumption']).replace(".",",")
    doc.tables[2].rows[3].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.tables[2].rows[2].cells[1].text = str(cblank['temperature']).replace(".",",")
    doc.tables[2].rows[2].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER


    #doc.tables[3].rows[0].cells[0].paragraphs[0].add_run().add_picture('C:\\Users\\kushhov\\Desktop\\vector-main\\scheme.jpg', width=Mm(90)).alignment =  WD_ALIGN_PARAGRAPH.CENTER #, width=Inches(5), height=Inches(3)
    #doc.paragraphs[5].add_run().add_picture('C:\\Users\\kushhov\\Desktop\\vector-main\\\Scheme\\scheme.jpg', width=Mm(90)).alignment =  WD_ALIGN_PARAGRAPH.CENTER #, width=Inches(5), height=Inches(3)


    path ='2-С.jpg'


    doc_scheme = f"{type_scheme[0]}-{type_scheme[1].upper()}"
  
    
    #st.write(doc_scheme)

    
    match doc_scheme:
        case "1-С":
            path = '1-С.jpg'
        case "2-С":
            path = '2-С.jpg'
        case "2-Ш":
            path = '2-Ш.jpg'
        case "3-С":
            path = '3-С.jpg'
        case "4-С":
            path = '4-С.jpg'
        case "4-Ш":
            path = '4-Ш.jpg'
        case "4М-С":
            path = '4-С.jpg'            
        case "5М-С":
            path = '5М-С.jpg'
        case "5М-Ш":
            path = '5М-Ш.jpg'
        case "5-С":
            path = '5-С.jpg'
        case "5-Ш":
            path = '5-Ш.jpg'
        case "6М-Ш":
            path = '6М-Ш.jpg'
        case "6-Ш":
            path = '6-Ш.jpg'

    #st.write(path)
    path ='Scheme/' + path
    #path = f"./{path}"
   

    doc.paragraphs[58].add_run().add_picture(path, width=Mm(180)).alignment =  WD_ALIGN_PARAGRAPH.CENTER


    doc.tables[2].autofit
    doc.tables[3].autofit
    #Data_frame = pd.read_excel('C:\\Users\\kushhov\\Desktop\\vector-main\\Table_vector.xlsx')
    #sheet_id = "1Qn-rGHE-mBaXHzVL9GHgGURV2OK0ZMRxcWerJ98qQhY"
    #Data_frame = pd.read_csv(f'https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv') неудачная попытка напрямую обрратиться к google таблице по ссылки
    
    #Data_frame = table_vector()

    #type_scheme = f"{type_scheme[0]}-{type_scheme[1]}-{type_scheme[2]}"
    #type_scheme = type_scheme.replace("С","C")
    Number_Vector = f"{type_scheme[0]}-{type_scheme[1]}-{type_scheme[2]}"
    Number_Vector = Number_Vector.upper().replace('С','C')
    K_Number_Vector = f"К-{Number_Vector}"
    for i in range(1,18):
        values_equipment = Data_frame[Number_Vector].values[i-1]
        count_equipment =  Data_frame[K_Number_Vector].values[i-1]
        doc.tables[3].rows[i].cells[1].text = str(values_equipment)
        doc.tables[3].rows[i].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[3].rows[i].cells[2].text = str(count_equipment)
        doc.tables[3].rows[i].cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.tables[3].autofit

    #doc.add_picture('C:\\Users\\kushhov\\Desktop\\vector-main\\scheme.jpg')
    return doc

   # file_zip.write('C:\\Users\\kushhov\\Desktop\\vector-main\\test'+str(x+1)+'.docx', compress_type=zipfile.ZIP_DEFLATED)
   
